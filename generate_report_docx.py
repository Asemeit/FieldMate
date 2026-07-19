"""Generate Word-ready FieldMate Chapters 7-9 aligned with user's Chapters 1-6."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = r"C:\Users\S A R D A R\Downloads\fieldwork\fieldwork\FieldMate_Chapters_7_to_9_v2.docx"


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    for level, size in [(1, 16), (2, 14), (3, 13)]:
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.size = Pt(size)


def add_center_title(doc: Document, text: str, size: int = 16) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_body(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    doc.add_paragraph()


def add_checkbox_line(doc: Document, label: str, options: list[str]) -> None:
    add_body(doc, label, bold=True)
    add_body(doc, "   ".join(f"[ ] {o}" for o in options))


def build() -> None:
    doc = Document()
    set_doc_defaults(doc)

    add_center_title(doc, "AI-ENABLED CROP DISEASE ADVISOR (FieldMate)", 18)
    add_center_title(doc, "Chapters 7, 8 and 9", 14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Student: [YOUR FULL NAME]  |  Reg. No.: [YOUR REG NO.]\n"
        "Institution: [YOUR UNIVERSITY]\n"
        "Department: [DEPARTMENT]\n"
        "Supervisor: [SUPERVISOR NAME]\n"
        "Study Area: Uasin Gishu County, Kenya\n"
        "Date: May 2026"
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    doc.add_page_break()

    # CHAPTER 7
    doc.add_heading("CHAPTER 7: IMPLEMENTATION (PROTOTYPE FRAMEWORK)", level=1)

    doc.add_heading("7.0 Introduction", level=2)
    add_body(
        doc,
        "This chapter presents the implementation of the AI-Enabled Crop Disease Advisor "
        "(prototype name: FieldMate), a Progressive Web Application (PWA) developed to address "
        "the crop disease management challenges identified in Chapters One and Five. As stated "
        "in Section 1.2, farmers in Kenya — including Uasin Gishu County — face late disease "
        "detection, limited access to extension officers, dependence on inaccurate manual methods, "
        "and high cost of modern diagnostic tools.",
    )
    add_body(
        doc,
        "Chapter Five confirmed these findings: 46.7% of respondents identified delayed disease "
        "detection as the major challenge, 36.7% reported inaccurate diagnosis, and 35% cited "
        "high cost of agricultural experts. At the same time, 73.3% expected early disease detection "
        "as the primary benefit of an AI advisory system, and 83.3% indicated they would recommend "
        "such a system.",
    )
    add_body(
        doc,
        "The prototype was implemented as an offline-first, voice-assisted PWA targeting "
        "smallholder farmers in Uasin Gishu County (Eldoret), supporting maize, potato, tomato, "
        "wheat, and beans as defined in Section 1.5. Implementation followed the three-layer "
        "architecture in Chapter Four (Presentation, Application, Database) and the testing plan "
        "in Chapter Three. The overall goal in Section 1.3.1 was pursued through a Minimum Viable "
        "Product (MVP) within the 14-week duration (Section 1.6).",
    )

    doc.add_heading("7.1 System Implementation", level=2)
    doc.add_heading("7.1.1 Implementation Approach and Link to Research Objectives", level=3)
    add_table(
        doc,
        ["Reference", "Implementation Response"],
        [
            ["1.3.3.1 / Ch. 2.2.1 — Traditional challenges", "Image upload + AI/ML pipeline; digital diagnosis history"],
            ["1.3.3.2 / Ch. 2.2.2 — AI benefits", "Early detection, voice guidance, offline PWA, weather alerts, PDF reports"],
            ["1.3.2 — AI-powered detection", "TensorFlow.js ML + optional Claude Vision + offline knowledge base"],
            ["1.3.2 — Voice-guided recommendations", "Web Speech API; English and Kiswahili; adjustable voice speed"],
            ["1.3.2 — Weather-integrated predictor", "Open-Meteo API; rule-based risk (humidity, temp, rainfall thresholds)"],
            ["1.3.2 — Global accessibility", "Mobile-first UI, large touch targets, onboarding tour"],
            ["4.4.1 — Three-layer architecture", "React UI, TypeScript services, IndexedDB database"],
            ["5.3–5.4 — User expectations", "Simple interface for users with limited smartphone-agriculture experience"],
        ],
    )

    doc.add_heading("7.1.2 Core Modules Implemented", level=3)
    add_bullets(doc, [
        "User Management: registration, login, logout, password reset, salted SHA-256 hashing, profile update (auth.ts).",
        "Disease Detection: camera/gallery upload, crop selection, TensorFlow.js ML (maize/potato/tomato), optional Claude Vision, offline demo database for wheat/beans (DetectPage, mlClassifier.ts, claude.ts).",
        "Voice Guidance: bilingual content, Text-to-Speech, voice speed control, dashboard/results play buttons (speech.ts).",
        "Weather Module: Open-Meteo for Eldoret, rule-based disease risk alerts, offline cache (weather.ts, WeatherPage).",
        "History & Reports: IndexedDB storage, history CRUD, PDF export via jsPDF (HistoryPage, pdfExport.ts).",
        "PWA: vite-plugin-pwa, service worker, installable on Android, offline asset caching.",
        "Accessibility: EN/SW toggle, interactive 6-step onboarding tour, Uasin Gishu pilot banner.",
    ])

    doc.add_heading("7.1.3 Prototype Scope vs. Project Objectives (Section 1.3.2)", level=3)
    add_table(
        doc,
        ["Objective Target", "Prototype Implementation"],
        [
            ["≥92% accuracy, 10+ crop-disease pairs", "MVP: 6 ML classes (PlantVillage) + offline entries for wheat/beans; confidence % displayed"],
            ["Nine language voice support", "English and Kiswahili implemented; others planned (Section 6.4)"],
            ["Weather risk (humidity >80%, 18–28°C)", "Rule-based engine implemented with Open-Meteo data"],
            ["Global accessibility", "Mobile-first PWA with voice, large buttons, high-contrast theme"],
        ],
    )

    doc.add_heading("7.1.4 Testing (Chapter 3.6)", level=3)
    add_table(
        doc,
        ["Test Type", "Application", "Result"],
        [
            ["Unit testing", "Auth, validation, weather rules", "Individual functions verified"],
            ["Integration testing", "Camera → ML → IndexedDB → History", "End-to-end scan workflow functional"],
            ["System testing", "Register, scan, PDF, weather, voice", "All modules working in Chrome"],
            ["UAT", "Questionnaire n=60 (Chapter 5)", "83.3% would recommend AI system"],
        ],
    )

    doc.add_heading("7.2 Technologies Used", level=2)

    doc.add_heading("7.2.1 Hardware Platform", level=3)
    add_table(
        doc,
        ["Component", "Specification", "Purpose"],
        [
            ["Developer laptop", "Windows 10/11, 8 GB+ RAM", "Development and testing"],
            ["Farmer smartphone", "Android 8.0+, 5 MP camera", "Primary end-user device (Section 1.5)"],
            ["Internet", "Wi-Fi / 3G/4G", "Weather API, optional Claude API, initial app load"],
        ],
    )

    doc.add_heading("7.2.2 Programming Language", level=3)
    add_table(
        doc,
        ["Language", "Version", "Usage"],
        [
            ["TypeScript", "6.0.x", "Primary application code"],
            ["JavaScript", "ES2022+", "Compiled output, service worker"],
            ["HTML5 / CSS3", "—", "Structure and Tailwind CSS styling"],
            ["JSON", "—", "ML labels, configuration"],
        ],
    )

    doc.add_heading("7.2.3 Programming Tools", level=3)
    add_bullets(doc, [
        "Node.js & npm — runtime and packages",
        "Vite 8 — dev server and production bundler",
        "Visual Studio Code / Cursor IDE — development",
        "Google Chrome DevTools — IndexedDB, mobile emulation",
        "ESLint & Git — quality and version control",
        "Google Colab — ML model experimentation (Section 1.6)",
    ])

    doc.add_heading("7.2.4 Software Platform", level=3)
    add_table(
        doc,
        ["Software", "Version", "Function"],
        [
            ["React", "19.2.x", "UI framework (Presentation Layer, Ch. 4.4.1)"],
            ["TensorFlow.js", "4.22.x", "On-device ML (PlantVillage model)"],
            ["IndexedDB", "Browser API", "FieldMateDB v2 — users, diagnoses, settings, weather"],
            ["jsPDF", "4.2.x", "PDF report generation"],
            ["vite-plugin-pwa", "1.3.x", "Offline PWA (Section 2.2.2, benefit 5)"],
            ["Open-Meteo API", "REST", "Weather data for Eldoret"],
            ["Web Speech API", "Browser-native", "Voice guidance (Section 1.3.2)"],
            ["Claude Vision API", "Optional", "Live AI analysis when configured"],
        ],
    )

    doc.add_heading("7.3 Features of the Prototype", level=2)
    doc.add_heading("7.3.1 Technical Manual Screenshots", level=3)
    tech = [
        "Figure 7.1 — VS Code project structure (src/pages, src/services).",
        "Figure 7.2 — Chrome DevTools → IndexedDB → FieldMateDB (Ch. 4.4.6 ERD).",
        "Figure 7.3 — Hashed passwordHash in users store.",
        "Figure 7.4 — TensorFlow.js model loading from /models/plant-disease/.",
        "Figure 7.5 — Results page with ML Model badge and confidence %.",
        "Figure 7.6 — Service files: mlClassifier.ts, db.ts, auth.ts.",
        "Figure 7.7 — vite.config.ts PWA configuration.",
        "Figure 7.8 — Successful npm run build output.",
        "Figure 7.9 — Analysis mode badges: ML / Live AI / Demo.",
        "Figure 7.10 — Mobile emulation showing phone-frame layout.",
    ]
    add_bullets(doc, tech)

    doc.add_heading("7.3.2 User Manual Screenshots — Main Activity Step by Step", level=3)
    steps = [
        "Step 1 / Fig. 7.11 — Open app landing page.",
        "Step 2 / Fig. 7.12 — Register farmer account.",
        "Step 3 / Fig. 7.13 — Log in.",
        "Step 4 / Fig. 7.14 — Complete onboarding tour.",
        "Step 5 / Fig. 7.15 — View dashboard (Uasin Gishu pilot).",
        "Step 6 / Fig. 7.16 — Check weather and disease risk.",
        "Step 7 / Fig. 7.17 — Select crop and open camera.",
        "Step 8 / Fig. 7.18 — Capture leaf with positioning guide.",
        "Step 9 / Fig. 7.19 — View AI diagnosis and treatment advice.",
        "Step 10 / Fig. 7.20 — Play voice guidance (EN/SW).",
        "Step 11 / Fig. 7.21 — Download PDF report.",
        "Step 12 / Fig. 7.22 — Review scan history.",
        "Step 13 / Fig. 7.23 — Update profile and settings.",
        "Step 14 / Fig. 7.24 — Switch to Kiswahili.",
        "Step 15 / Fig. 7.25 — Log out.",
    ]
    add_bullets(doc, steps)
    add_body(doc, "[ INSERT SCREENSHOT HERE ] — paste images under each figure caption in Word.")

    doc.add_heading("7.4 Database Management System", level=2)
    add_body(
        doc,
        "The prototype implements Chapter Four Section 4.4.6 (ERD) using IndexedDB (FieldMateDB, "
        "version 2) in src/services/db.ts. IndexedDB supports offline-first operation (Section 2.2.2), "
        "requires no separate server, and provides CRUD operations (Section 4.1.1).",
    )
    add_table(
        doc,
        ["Object Store", "Primary Key", "Main Fields"],
        [
            ["users", "email", "name, passwordHash, county, createdAt"],
            ["diagnoses", "id", "cropType, diseaseName, confidence, imageUrl, recommendation, timestamp, syncStatus, analysisMode"],
            ["settings", "key", "value (language, voiceSpeed, activeUser)"],
            ["weather", "location", "temperature, humidity, rainfall, riskLevel, riskAlerts, timestamp"],
        ],
    )
    add_body(doc, "Normalisation (3NF): users stored once; settings as key-value; diagnoses as atomic records; weather cached by location.")
    add_table(
        doc,
        ["Entity", "Create", "Read", "Update", "Delete"],
        [
            ["users", "Register", "Login", "Update profile", "—"],
            ["diagnoses", "Save scan", "History, Results", "—", "Delete / Clear all"],
            ["settings", "saveSetting", "getSetting", "updateSettings", "—"],
            ["weather", "saveWeather", "getCachedWeather", "—", "—"],
        ],
    )

    doc.add_heading("7.4.1 Implementation vs. Chapter Five Findings", level=3)
    add_table(
        doc,
        ["Finding (Ch. 5)", "System Response"],
        [
            ["46.7% — delayed detection", "Instant ML/AI analysis; results in seconds"],
            ["35% — poor record management", "IndexedDB history + PDF export"],
            ["35% — high expert cost", "Self-service diagnosis on farmer's phone"],
            ["73.3% — early detection benefit", "Image-based AI with confidence %"],
            ["83.3% — would recommend", "Working MVP demonstrated to respondents"],
        ],
    )

    doc.add_heading("7.5 Chapter Summary", level=2)
    add_body(
        doc,
        "This chapter demonstrated implementation of the AI-Enabled Crop Disease Advisor as the "
        "FieldMate PWA, translating Chapters One, Four, and Five into a working system. Core features "
        "— AI/ML detection, voice guidance, weather alerts, offline operation, PDF reports, and digital "
        "history — address Section 2.2.1 challenges and deliver Section 2.2.2 benefits. Future work "
        "(Section 6.4) includes additional languages, expanded crop models, and wider field validation.",
    )

    doc.add_page_break()

    # CHAPTER 8
    doc.add_heading("CHAPTER 8: REFERENCES", level=1)
    refs = [
        "Abdullahi, A. S., Mahmud, M. S., & Alkali, A. M. (2022). Mobile-based plant disease detection using deep learning: A review. Computers and Electronics in Agriculture, 200, 107194. https://doi.org/10.1016/j.compag.2022.107194",
        "Food and Agriculture Organization. (2021). Integrated pest management: An introduction. FAO. https://www.fao.org",
        "Kenya Agricultural and Livestock Research Organization. (2023). Digital agricultural services in Kenya. KALRO. https://www.kalro.org",
        "Mohanty, S. P., Hughes, D. P., & Salathé, M. (2016). Using deep learning for image-based plant disease detection. Frontiers in Plant Science, 7, 1419. https://doi.org/10.3389/fpls.2016.01419",
        "Mozilla Developer Network. (2024). Progressive web apps. https://developer.mozilla.org/en-US/docs/Web/Progressive_web_apps",
        "Open-Meteo. (2024). Weather API documentation. https://open-meteo.com/en/docs",
        "Penn State University. (2024). PlantVillage. https://plantvillage.psu.edu/",
        "Republic of Kenya. (2018). National climate smart agriculture strategy 2017–2026. Ministry of Agriculture, Livestock, Fisheries and Irrigation.",
        "TensorFlow Team. (2024). TensorFlow.js. https://www.tensorflow.org/js",
        "United Nations. (2015). Transforming our world: The 2030 Agenda for Sustainable Development (SDG 2: Zero Hunger). https://sdgs.un.org/goals",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)

    doc.add_page_break()

    # CHAPTER 9
    doc.add_heading("CHAPTER 9: APPENDIX", level=1)
    doc.add_heading("9.1 Questionnaire / Interview Schedule", level=2)
    add_body(doc, "Instrument used for Chapter Five data collection (n=60 responses). Study area: Uasin Gishu County.")
    add_table(
        doc,
        ["Section", "Questions", "Maps To"],
        [
            ["A", "A1–A10", "Demographics (Section 5.2)"],
            ["B", "B1–B5", "Objective 1.3.3.1 / Section 2.2.1 — Traditional challenges"],
            ["B", "B6–B8", "Objective 1.3.3.2 / Section 2.2.2 — AI benefits"],
            ["B", "B9–B11", "Section 1.3.2 — Voice, Kiswahili, features"],
            ["B", "B12–B14", "Section 5.4 — Acceptance and recommendation"],
        ],
    )

    doc.add_heading("SECTION A: Respondent Biodata", level=3)
    add_checkbox_line(doc, "A1. Gender:", ["Male", "Female", "Prefer not to say"])
    add_checkbox_line(doc, "A2. Age group:", ["Below 18", "18–25", "26–35", "36–45", "46–55", "56+"])
    add_checkbox_line(doc, "A3. Education level:", ["Primary", "Secondary", "Certificate/Diploma", "University", "Postgraduate"])
    add_checkbox_line(doc, "A4. Category:", ["Student/Researcher", "Smallholder Farmer", "Farm Worker", "Extension Officer", "Other"])
    add_checkbox_line(doc, "A5. Years in agriculture:", ["<1 year", "1–3 years", "4–10 years", ">10 years"])
    add_body(doc, "A6. Crops (tick all): [ ] Maize  [ ] Wheat  [ ] Potatoes  [ ] Beans  [ ] Tomatoes")
    add_checkbox_line(doc, "A7. Farm size:", ["<1 acre", "1–3 acres", "4–10 acres", ">10 acres"])
    add_checkbox_line(doc, "A8. Smartphone:", ["Yes — personal", "Yes — shared", "No"])
    add_checkbox_line(doc, "A9. Internet:", ["Reliable", "Sometimes", "Rarely", "Never"])
    add_checkbox_line(doc, "A10. Language:", ["English", "Kiswahili", "Both"])

    doc.add_heading("SECTION B: Crop Disease Management", level=3)
    add_checkbox_line(doc, "B1. Disease outbreaks per season:", ["Never", "Once", "2–3 times", "More than 3"])
    add_body(doc, "B2. Identification method (tick all): [ ] Manual  [ ] Expert  [ ] Lab  [ ] ML/App  [ ] Other")
    add_checkbox_line(doc, "B3. Time to get advice:", ["Same day", "1–3 days", "4–7 days", ">1 week", "Rarely"])
    add_body(doc, "B4. Difficulty without expert: [ ] 1  [ ] 2  [ ] 3  [ ] 4  [ ] 5")
    add_body(doc, "B5. Challenges (Ch. 2.2.1 — tick all): [ ] Delayed detection  [ ] Inaccurate diagnosis  [ ] High cost  [ ] Poor records  [ ] No monitoring  [ ] Limited extension  [ ] Literacy  [ ] Internet")
    add_body(doc, "B6. Expected AI benefits (Ch. 2.2.2 — tick all): [ ] Early detection  [ ] Accuracy  [ ] Reduced losses  [ ] Voice  [ ] Offline  [ ] Cost reduction  [ ] Records  [ ] Weather alerts")
    add_checkbox_line(doc, "B7. Trust confidence % shown?", ["Yes", "Yes with expert", "Unsure", "No"])
    add_checkbox_line(doc, "B8. Offline importance:", ["Very important", "Important", "Neutral", "Not important"])
    add_checkbox_line(doc, "B9. Voice importance:", ["Very important", "Important", "Neutral", "Not important"])
    add_checkbox_line(doc, "B10. Kiswahili importance:", ["Very important", "Important", "Neutral", "Not important"])
    add_body(doc, "B11. Rank features 1–5: Scan [ ] Weather [ ] Treatment [ ] Voice [ ] PDF [ ] History [ ]")
    add_checkbox_line(doc, "B12. Recommend AI advisor?", ["Yes", "Maybe", "No"])
    add_body(doc, "B13. Barriers: [ ] No phone  [ ] Poor internet  [ ] Literacy  [ ] Trust  [ ] Data cost  [ ] Training")
    add_body(doc, "B14. Ease of use after demo: [ ] 1  [ ] 2  [ ] 3  [ ] 4  [ ] 5")
    add_body(doc, "B15. Open-ended challenge: _______________________________________________")
    add_body(doc, "B16. Open-ended improvement: _______________________________________________")

    doc.add_heading("9.2 Work Plan (Gantt Chart) — 14 Weeks", level=2)
    weeks = [f"W{i}" for i in range(1, 15)]
    gantt = [
        ["Literature review (Ch. 2)"] + ["X", "X"] + [""] * 12,
        ["Questionnaire & Ch. 5 analysis"] + ["", "X", "X", "X"] + [""] * 10,
        ["System analysis & design (Ch. 4)"] + [""] * 2 + ["X", "X"] + [""] * 10,
        ["UI/UX design"] + [""] * 3 + ["X", "X"] + [""] * 9,
        ["IndexedDB database"] + [""] * 4 + ["X", "X"] + [""] * 8,
        ["Authentication module"] + [""] * 5 + ["X", "X"] + [""] * 7,
        ["Dashboard & navigation"] + [""] * 6 + ["X", "X"] + [""] * 6,
        ["Camera & image upload"] + [""] * 7 + ["X", "X"] + [""] * 5,
        ["TensorFlow.js ML"] + [""] * 8 + ["X", "X"] + [""] * 4,
        ["Weather (Open-Meteo)"] + [""] * 9 + ["X"] + [""] * 4,
        ["Voice & Kiswahili"] + [""] * 10 + ["X", "X"] + [""] * 2,
        ["PDF & history CRUD"] + [""] * 11 + ["X"] + [""] * 2,
        ["PWA & offline testing"] + [""] * 12 + ["X"] + [""],
        ["Report Ch. 6–7 & presentation"] + [""] * 12 + ["X", "X"],
    ]
    add_table(doc, ["Task"] + weeks, gantt)

    doc.add_heading("9.3 Budget (KES)", level=2)
    add_table(
        doc,
        ["Item", "Description", "Qty", "Unit", "Total"],
        [
            ["A1", "Android test smartphone", "1", "18,000", "18,000"],
            ["A2", "Phone stand & cable", "1", "1,500", "1,500"],
            ["B1", "Internet (14 weeks)", "14", "500", "7,000"],
            ["B2", "Questionnaires (60)", "60", "20", "1,200"],
            ["C1", "Transport Uasin Gishu", "6", "800", "4,800"],
            ["C2", "Focus group refreshments", "10", "200", "2,000"],
            ["D1", "Report binding (3 copies)", "3", "800", "2,400"],
            ["D2", "Presentation materials", "1", "1,500", "1,500"],
            ["E1", "Contingency 10%", "1", "3,840", "3,840"],
            ["", "GRAND TOTAL", "", "", "42,240"],
        ],
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
